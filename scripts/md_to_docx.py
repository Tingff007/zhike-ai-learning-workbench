# -*- coding: utf-8 -*-
"""Markdown 转 Word (.docx) 通用转换脚本

用途：把项目内中文技术文档（docs/*.md）转换为排版规范的 Word 文档。
支持元素：标题(1-3级)、表格、围栏代码块、引用块、无序列表、分隔线、
行内加粗与行内代码。输出中文字体（正文宋体、标题微软雅黑、代码 Consolas）。
用法：python scripts/md_to_docx.py <输入.md> <输出.docx>
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def set_run_font(run, name_ascii, name_east, size, bold=False, color=None, italic=False):
    """统一设置 run 的西文与中文（eastAsia）字体、字号、粗细与颜色"""
    run.font.name = name_ascii
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    r = run._element.rPr
    if r is None:
        r = run._element.get_or_add_rPr()
    rFonts = r.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = r.makeelement(qn("w:rFonts"), {})
        r.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name_east)


def add_inline_runs(paragraph, text, base_size, bold_all=False, mono=False):
    """按行内标记拆分文本并添加 run：**加粗**、`行内代码`"""
    tokens = re.split(r"(\*\*.+?\*\*|`[^`]+`)", text)
    for tok in tokens:
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            run = paragraph.add_run(tok[2:-2])
            set_run_font(run, "微软雅黑" if bold_all else "宋体", "微软雅黑" if bold_all else "宋体",
                         base_size, bold=True)
        elif tok.startswith("`") and tok.endswith("`") and len(tok) > 2:
            run = paragraph.add_run(tok[1:-1])
            set_run_font(run, "Consolas", "宋体", base_size - 0.5)
        else:
            run = paragraph.add_run(tok)
            set_run_font(run, "微软雅黑" if bold_all else "宋体", "微软雅黑" if bold_all else "宋体",
                         base_size, bold=bold_all)


def add_code_block(doc, code_text):
    """代码块：等宽字体 + 浅灰底纹 + 缩进"""
    for line in code_text.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(12)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(line if line else " ")
        set_run_font(run, "Consolas", "宋体", 9)
        # 浅灰底纹
        pPr = p._p.get_or_add_pPr()
        shd = pPr.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:fill"): "F2F2F2"})
        pPr.append(shd)


def add_table(doc, rows):
    """Markdown 表格：首行加粗，Table Grid 边框"""
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(n_cols):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            cell_text = row[j] if j < len(row) else ""
            add_inline_runs(p, cell_text.strip(), 9, bold_all=(i == 0))
    doc.add_paragraph()  # 表后空行


def convert(md_path, docx_path):
    text = Path(md_path).read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    # 页边距与默认样式
    for section in doc.sections:
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)

    # 标题样式中文字体（覆盖内置 Heading 样式的主题字体）
    for style_name, size in (("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)):
        st = doc.styles[style_name]
        st.font.name = "微软雅黑"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1F, 0x3B, 0x66)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    i = 0
    n = len(lines)
    in_code = False
    code_buf = []
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # 围栏代码块
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                in_code = False
                add_code_block(doc, "\n".join(code_buf))
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # 分隔线
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # 标题
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            title_text = re.sub(r"\*\*(.+?)\*\*", r"\1", m.group(2))
            h = doc.add_heading("", level=min(level, 3))
            run = h.add_run(title_text)
            set_run_font(run, "微软雅黑", "微软雅黑", {1: 18, 2: 14, 3: 12}.get(level, 12), bold=True)
            if level == 1:
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # 表格：当前行含 | 且下一行为分隔行
        if stripped.startswith("|") and stripped.endswith("|") and i + 1 < n:
            next_stripped = lines[i + 1].strip()
            if re.match(r"^\|[\s:\-|]+\|$", next_stripped):
                rows = []
                rows.append([c.strip() for c in stripped.strip("|").split("|")])
                i += 2
                while i < n and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                    rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                    i += 1
                add_table(doc, rows)
                continue

        # 引用块
        if stripped.startswith(">"):
            quote_text = stripped.lstrip(">").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            run = p.add_run(quote_text)
            set_run_font(run, "宋体", "宋体", 10.5, italic=True, color=RGBColor(0x59, 0x59, 0x59))
            i += 1
            continue

        # 无序列表
        if re.match(r"^[-*]\s+", stripped):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, re.sub(r"^[-*]\s+", "", stripped), 10.5)
            i += 1
            continue

        # 普通段落（含空行）
        if stripped:
            p = doc.add_paragraph()
            add_inline_runs(p, stripped, 10.5)
        i += 1

    doc.save(docx_path)
    print(f"OK -> {docx_path}")


if __name__ == "__main__":
    src = sys.argv[1]
    dst = sys.argv[2]
    convert(src, dst)
